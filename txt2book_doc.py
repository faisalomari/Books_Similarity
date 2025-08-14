import os
import requests
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def reverse_parentheses_in_text(text):
    # This function reverses the parentheses around footnote numbers.
    return re.sub(r'(\(\d+\))', lambda match: f"({match.group(0)[1:-1]})", text)

def set_page_settings(doc):
    # Set page margins to narrow to help fit the content in one page
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(30)  # Left margin in points
        section.right_margin = Pt(30)  # Right margin in points
        section.top_margin = Pt(30)  # Top margin in points
        section.bottom_margin = Pt(30)  # Bottom margin in points

def set_font_for_paragraph(paragraph):
    # Set the font size and family for the paragraph
    run = paragraph.runs[0]  # Get the first run of the paragraph
    run.font.size = Pt(10)  # Set font size to fit text
    run.font.name = 'Times New Roman'  # Set font family (use a standard one)

def txt_to_docx(txt_filename, doc):
    # Read the content of the txt file
    with open(txt_filename, 'r', encoding='utf-8') as file:
        text = file.read()

    # Split the content into main content and footnotes (حاشية)
    if "(حاشية)" in text:
        main_content, footnotes_section = text.split("(حاشية)")
    else:
        main_content = text
        footnotes_section = ""

    # Reverse parentheses around footnote numbers in the main content
    main_content = reverse_parentheses_in_text(main_content)

    # Add main content to the Word document with RTL alignment for Arabic text
    paragraph = doc.add_paragraph(main_content.strip())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align the main content to the right for RTL
    set_font_for_paragraph(paragraph)  # Set font size and family

    # If there are footnotes (حاشية), add them at the end
    if footnotes_section.strip():
        footnotes_paragraph = doc.add_paragraph("\n(حاشية)")
        footnotes_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align the footnotes to the right for RTL
        set_font_for_paragraph(footnotes_paragraph)  # Set font size and family

        # Split the footnotes by lines
        footnote_lines = footnotes_section.strip().split("\n")
        
        for line in footnote_lines:
            # Reverse parentheses in footnote content
            line = reverse_parentheses_in_text(line.strip())
            # Add each footnote on a new line with RTL alignment
            footnote_paragraph = doc.add_paragraph(line)
            footnote_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_font_for_paragraph(footnote_paragraph)  # Set font size and family

        # Add a page break to move to the next page after footnotes
        doc.add_paragraph().add_run().add_break()

    # Add a page break after each .txt file content to start a new page
    doc.add_paragraph().add_run().add_break()

def process_all_txt_files_in_folder(folder_path, docx_filename):
    # Create a new Word document
    doc = Document()

    # Set page settings to help fit content in one page
    set_page_settings(doc)

    # Sort the txt files by their name to preserve the order
    txt_files = sorted([f for f in os.listdir(folder_path) if f.endswith(".txt")])

    # Loop through all the txt files in the folder in sorted order
    for txt_filename in txt_files:
        full_txt_filename = os.path.join(folder_path, txt_filename)
        print(f"Processing {full_txt_filename}...")
        txt_to_docx(full_txt_filename, doc)

    # Save the document to a .docx file
    doc.save(docx_filename)
    print(f"All .txt files have been converted to {docx_filename}")

# Example usage
folder_path = "book"  # Path to the folder containing your .txt files
docx_filename = "combined_book.docx"  # Output .docx file path

# Process all .txt files in the folder and save to a single .docx file
process_all_txt_files_in_folder(folder_path, docx_filename)
