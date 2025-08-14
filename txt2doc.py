import requests
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

def reverse_parentheses_in_text(text):
    # This function reverses the parentheses around footnote numbers.
    return re.sub(r'(\(\d+\))', lambda match: f"({match.group(0)[1:-1]})", text)

def set_page_settings(doc):
    # Set page margins to narrow to help fit the content in one page
    sections = doc.sections
    for section in sections:
        section.left_margin = Pt(30)  # Left margin in points
        section.right_margin = Pt(30)  # Right margin in points
        section.top_margin = Pt(30)  # Top margin in points
        section.bottom_margin = Pt(30)  # Bottom margin in points

def set_font_for_paragraph(paragraph):
    # Set the font size and family for the paragraph
    run = paragraph.runs[0]  # Get the first run of the paragraph
    run.font.size = Pt(10)  # Set font size to fit text
    run.font.name = 'Times New Roman'  # Set font family (use a standard one)

def txt_to_docx(txt_filename, docx_filename):
    # Read the content of the txt file
    with open(txt_filename, 'r', encoding='utf-8') as file:
        text = file.read()

    # Create a new Word document
    doc = Document()

    # Set page settings to help fit content in one page
    set_page_settings(doc)

    # Split the content into main content and footnotes (حاشية)
    if "(حاشية)" in text:
        main_content, footnotes_section = text.split("(حاشية)")
    else:
        main_content = text
        footnotes_section = ""

    # Reverse parentheses around footnote numbers in the main content
    main_content = reverse_parentheses_in_text(main_content)

    # Add main content to the Word document with RTL alignment for Arabic text
    paragraph = doc.add_paragraph(main_content.strip())
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align the main content to the right for RTL
    set_font_for_paragraph(paragraph)  # Set font size and family

    # If there are footnotes (حاشية), add them at the end
    if footnotes_section.strip():
        footnotes_paragraph = doc.add_paragraph("\n(حاشية)")
        footnotes_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT  # Align the footnotes to the right for RTL
        set_font_for_paragraph(footnotes_paragraph)  # Set font size and family

        # Split the footnotes by lines
        footnote_lines = footnotes_section.strip().split("\n")
        
        for line in footnote_lines:
            # Reverse parentheses in footnote content
            line = reverse_parentheses_in_text(line.strip())
            # Add each footnote on a new line with RTL alignment
            footnote_paragraph = doc.add_paragraph(line)
            footnote_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            set_font_for_paragraph(footnote_paragraph)  # Set font size and family

    # Save the document to a .docx file
    doc.save(docx_filename)

# Example usage
name = "book_page_5"
txt_filename = f"book/{name}.txt"  # Replace with your .txt file path
docx_filename = f"{name}.docx"  # Output .docx file path

# Convert the .txt file to .docx
txt_to_docx(txt_filename, docx_filename)
print(f"Converted {txt_filename} to {docx_filename}")
