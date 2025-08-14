# shamela_tools.py
# -*- coding: utf-8 -*-
import os, re, csv, time, glob
from typing import List, Dict, Tuple, Optional
import requests
from bs4 import BeautifulSoup
from tqdm import tqdm
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import difflib

# ---------------------------
# HTTP helpers
# ---------------------------
DEFAULT_HEADERS = {
    "User-Agent": "Mozilla/5.0 (X11; Linux x86_64) AppleWebKit/537.36 "
                  "(KHTML, like Gecko) Chrome/119.0.0.0 Safari/537.36"
}

def fetch_shamela_page_html(book_id: int, page_num: int, timeout: int = 20) -> Optional[str]:
    """
    Returns HTML text for a given shamela book page or None if not found.
    """
    url = f"https://shamela.ws/book/{book_id}/{page_num}#p{page_num}"
    try:
        r = requests.get(url, headers=DEFAULT_HEADERS, timeout=timeout)
        if r.status_code != 200:
            return None
        return r.text
    except requests.RequestException:
        return None

def extract_text_from_shamela_html(html: str) -> Optional[str]:
    """
    Extracts the main text including inline footnote numbers and appends (حاشية) if exists.
    """
    soup = BeautifulSoup(html, 'html.parser')
    content = soup.find('div', class_='nass margin-top-10')
    if not content:
        return None

    paragraphs = content.find_all('p')
    text_content = ""
    footnotes = []

    for p in paragraphs:
        # collect footnotes (حاشية)
        if 'hamesh' in p.get('class', []):
            footnote_text = p.get_text(separator="\n", strip=True)
            footnotes.append(footnote_text)
            continue

        # drop hyperlinks
        for a in p.find_all('a'):
            a.decompose()

        # keep footnote numbers inline
        paragraph_text = p.get_text(strip=True)
        footnote_numbers = p.find_all('span', class_='c2')
        for fn in footnote_numbers:
            n = fn.get_text(strip=True)
            paragraph_text = paragraph_text.replace(f"({n})", f" ({n})")

        if paragraph_text:
            text_content += paragraph_text + "\n\n"

    if footnotes:
        text_content += "\n\n(حاشية)\n"
        for footnote in footnotes:
            text_content += footnote + "\n"

    return text_content.strip() or None

# ---------------------------
# Extraction to per-page TXT
# ---------------------------
def save_text(path: str, text: str) -> None:
    os.makedirs(os.path.dirname(path) or ".", exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        f.write(text)

def extract_book_to_txt(book_id: int,
                        out_dir: str,
                        start_page: int = 1,
                        max_pages: int = 5000,
                        consecutive_empty_to_stop: int = 2,
                        sleep_between: float = 0.0) -> int:
    """
    Sequentially fetch pages and save as:
      out_dir/book_page_{i}.txt
    Stops when `consecutive_empty_to_stop` pages in a row have no content or we reach `max_pages`.
    Returns count of pages saved.
    """
    os.makedirs(out_dir, exist_ok=True)
    saved = 0
    empty_streak = 0

    for p in tqdm(range(start_page, start_page + max_pages), desc=f"Book {book_id}: extracting pages"):
        html = fetch_shamela_page_html(book_id, p)
        if not html:
            empty_streak += 1
        else:
            text = extract_text_from_shamela_html(html)
            if not text:
                empty_streak += 1
            else:
                empty_streak = 0
                save_text(os.path.join(out_dir, f"book_page_{p}.txt"), text)
                saved += 1

        if empty_streak >= consecutive_empty_to_stop:
            break

        if sleep_between > 0:
            time.sleep(sleep_between)

    return saved

# ---------------------------
# DOCX building
# ---------------------------
def _reverse_paren_numbers(text: str) -> str:
    # keep “(123)” shape; hook here if you ever need to flip RTL parens
    return re.sub(r'(\(\d+\))', lambda m: f"({m.group(0)[1:-1]})", text)

def _set_page_margins(doc: Document, pt: int = 30):
    for s in doc.sections:
        s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Pt(pt)

def _set_para_font(paragraph):
    if not paragraph.runs:
        paragraph.add_run()
    run = paragraph.runs[0]
    run.font.size = Pt(10)
    run.font.name = 'Times New Roman'  # use an available font in your system

def combine_txt_folder_to_docx(txt_dir: str, docx_path: str) -> None:
    """
    Reads all files like book_page_*.txt (sorted by page number) and writes a single RTL .docx.
    """
    doc = Document()
    _set_page_margins(doc, 30)

    files = sorted(
        glob.glob(os.path.join(txt_dir, "book_page_*.txt")),
        key=lambda p: int(re.search(r"book_page_(\d+)\.txt", os.path.basename(p)).group(1))
    )

    for fp in tqdm(files, desc=f"Building DOCX from {txt_dir}"):
        with open(fp, "r", encoding="utf-8") as f:
            text = f.read()

        # split main content / (حاشية)
        if "(حاشية)" in text:
            main_content, footnotes_section = text.split("(حاشية)", 1)
        else:
            main_content, footnotes_section = text, ""

        main_content = _reverse_paren_numbers(main_content)
        para = doc.add_paragraph(main_content.strip())
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_para_font(para)

        if footnotes_section.strip():
            ptitle = doc.add_paragraph("\n(حاشية)")
            ptitle.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            _set_para_font(ptitle)

            for line in footnotes_section.strip().splitlines():
                line = _reverse_paren_numbers(line.strip())
                p = doc.add_paragraph(line)
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                _set_para_font(p)

        # page break after each source page
        doc.add_paragraph().add_run().add_break()

    os.makedirs(os.path.dirname(docx_path) or ".", exist_ok=True)
    doc.save(docx_path)

# ---------------------------
# Normalization + sentence extraction from DOCX
# ---------------------------
ARABIC_DIAC = r"[\u0610-\u061A\u064B-\u065F\u0670\u06D6-\u06ED]"
PUNCT = r"[^\w\s\u0600-\u06FF]"   # keep Arabic letters/digits/underscore/space
SENT_SPLIT = re.compile(r"[\.!\?؟…]+|\n+")

def normalize_ar(text: str) -> str:
    if not text:
        return ""
    t = re.sub(ARABIC_DIAC, "", text)
    t = t.replace("ـ", "")
    t = t.replace("أ", "ا").replace("إ", "ا").replace("آ", "ا")
    t = t.replace("ى", "ي")
    t = t.replace("ة", "ه")
    t = re.sub(PUNCT, " ", t)
    t = re.sub(r"\s+", " ", t).strip()
    return t

def iter_docx_sentences(docx_path: str) -> List[Dict]:
    """
    Returns list of dicts: {"raw": str, "norm": str, "pos": "page:X"}
    Page counting uses manual page breaks we inserted.
    """
    doc = Document(docx_path)
    page_no = 1
    results = []
    para_idx = 0

    from docx.oxml.ns import qn
    for para in tqdm(doc.paragraphs, desc=f"Extracting from {os.path.basename(docx_path)}"):
        # detect page breaks
        for run in para.runs:
            for child in run._r:
                if child.tag == qn('w:br') and child.get(qn('w:type')) == 'page':
                    page_no += 1

        raw_para = para.text.strip()
        if not raw_para:
            para_idx += 1
            continue

        parts = [s.strip() for s in SENT_SPLIT.split(raw_para) if s.strip()]
        for s in parts:
            if len(s) < 12:
                continue
            results.append({
                "raw": s,
                "norm": normalize_ar(s),
                "pos": f"page:{page_no}" if page_no >= 1 else f"p:{para_idx}"
            })
        para_idx += 1
    return results

# ---------------------------
# Matching utilities
# ---------------------------
def overlap_chars(a: str, b: str) -> int:
    sm = difflib.SequenceMatcher(None, a, b, autojunk=False)
    return sum(block.size for block in sm.get_matching_blocks())

def match_books_sequential(book1_docx: str,
                           book2_docx: str,
                           out_csv: str,
                           min_len_chars: int = 24,
                           score_threshold: int = 60) -> None:
    """
    Compare sentences from book1 to book2 using RapidFuzz and write CSV.
    """
    from rapidfuzz import process, fuzz

    book1_sents = iter_docx_sentences(book1_docx)
    book2_sents = iter_docx_sentences(book2_docx)
    book2_norms = [s["norm"] for s in book2_sents]

    rows: List[Tuple[str, str, str, str, int, int]] = []

    for s1 in tqdm(book1_sents, desc="Matching sequentially"):
        if len(s1["raw"]) < min_len_chars:
            continue
        if len(s1["norm"].split()) < 4:
            continue

        m = process.extractOne(
            s1["norm"],
            book2_norms,
            scorer=fuzz.token_set_ratio,
            score_cutoff=score_threshold
        )
        if not m:
            continue

        _, score, idx2 = m
        strict_score = fuzz.QRatio(s1["norm"], book2_sents[idx2]["norm"])
        if strict_score < score_threshold - 5:
            continue

        pos1 = s1["pos"]
        pos2 = book2_sents[idx2]["pos"]
        raw1 = s1["raw"]
        raw2 = book2_sents[idx2]["raw"]
        ov   = overlap_chars(s1["norm"], book2_sents[idx2]["norm"])
        rows.append((pos1, pos2, raw1, raw2, int(score), ov))

    rows.sort(key=lambda r: (r[4], r[5]), reverse=True)

    os.makedirs(os.path.dirname(out_csv) or ".", exist_ok=True)
    with open(out_csv, "w", encoding="utf-8-sig", newline="") as f:
        w = csv.writer(f)
        w.writerow([
            "pos_in_book1",
            "pos_in_book2",
            "sentence_book1",
            "sentence_book2",
            "score",
            "overlap_chars"
        ])
        w.writerows(rows)

# ---------------------------
# Map CSV matches back to per-page TXT numbers
# ---------------------------
def load_pages(dir_path: str, pattern: str = "book_page_*.txt") -> List[Dict]:
    files = sorted(
        glob.glob(os.path.join(dir_path, pattern)),
        key=lambda p: int(re.search(r"book_page_(\d+)\.txt", os.path.basename(p)).group(1))
    )
    pages = []
    for fp in tqdm(files, desc=f"Loading pages from {dir_path}"):
        m = re.search(r"book_page_(\d+)\.txt", os.path.basename(fp))
        if not m:
            continue
        page_no = int(m.group(1))
        with open(fp, "r", encoding="utf-8") as f:
            raw = f.read()
        pages.append({
            "page": page_no,
            "raw": raw,
            "norm": normalize_ar(raw),
            "path": fp
        })
    return pages

def find_page_for_sentence(sentence: str,
                           pages: List[Dict],
                           use_fuzzy: bool = False,
                           fuzzy_cutoff: int = 92) -> Optional[int]:
    s_norm = normalize_ar(sentence)
    if not s_norm:
        return None

    # exact normalized substring search
    for p in pages:
        if s_norm in p["norm"]:
            return p["page"]

    if not use_fuzzy:
        return None

    try:
        from rapidfuzz import fuzz
    except ImportError:
        return None

    best_page, best_score = None, -1
    # conservative partial matching
    for p in pages:
        score = fuzz.partial_ratio(s_norm, p["norm"])
        if score > best_score:
            best_score, best_page = score, p["page"]

    return best_page if best_score >= fuzzy_cutoff else None

def update_csv_with_pages(csv_in: str,
                          csv_out: str,
                          book1_pages_dir: str,
                          book2_pages_dir: str,
                          book1_use_fuzzy: bool = False,
                          book2_use_fuzzy: bool = False,
                          cols=("pos_in_book1","pos_in_book2","sentence_book1","sentence_book2")) -> None:
    b1_pages = load_pages(book1_pages_dir)
    b2_pages = load_pages(book2_pages_dir)

    with open(csv_in, "r", encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        fieldnames = reader.fieldnames or []
        rows_in = list(reader)

    needed = set(cols)
    if not needed.issubset(set(fieldnames)):
        raise RuntimeError(f"CSV missing required columns: {needed - set(fieldnames)}")

    rows_out: List[Dict] = []
    for r in tqdm(rows_in, desc="Assigning page numbers"):
        s1 = r.get("sentence_book1", "") or r.get("sentence_ihya", "") or ""
        s2 = r.get("sentence_book2", "") or r.get("sentence_qut",  "") or ""

        p1 = find_page_for_sentence(s1, b1_pages, use_fuzzy=book1_use_fuzzy)
        p2 = find_page_for_sentence(s2, b2_pages, use_fuzzy=book2_use_fuzzy)

        if p1 is not None:
            r["pos_in_book1"] = str(p1)
        if p2 is not None:
            r["pos_in_book2"] = str(p2)

        rows_out.append(r)

    with open(csv_out, "w", encoding="utf-8-sig", newline="") as f:
        writer = csv.DictWriter(f, fieldnames=fieldnames)
        writer.writeheader()
        writer.writerows(rows_out)

# ---------------------------
# One-call pipeline
# ---------------------------
def run_full_pipeline(book1_id: int,
                      book2_id: int,
                      workdir: str = "output",
                      max_pages: int = 5000,
                      stop_after_empty: int = 2,
                      score_threshold: int = 60,
                      min_len_chars: int = 24,
                      use_fuzzy_backmap: bool = False) -> Dict[str, str]:
    """
    End-to-end:
      - scrape both books into per-page txt
      - build docx for each
      - match sentences and write CSV
      - update CSV with page numbers derived from txt pages
    Returns dict of produced file paths.
    """
    os.makedirs(workdir, exist_ok=True)
    b1_dir = os.path.join(workdir, f"book_{book1_id}", "book")
    b2_dir = os.path.join(workdir, f"book_{book2_id}", "book")

    # 1) Extract pages
    extract_book_to_txt(book1_id, b1_dir, max_pages=max_pages, consecutive_empty_to_stop=stop_after_empty)
    extract_book_to_txt(book2_id, b2_dir, max_pages=max_pages, consecutive_empty_to_stop=stop_after_empty)

    # 2) DOCX build
    b1_docx = os.path.join(workdir, f"book_{book1_id}", "combined_book.docx")
    b2_docx = os.path.join(workdir, f"book_{book2_id}", "combined_book.docx")
    combine_txt_folder_to_docx(b1_dir, b1_docx)
    combine_txt_folder_to_docx(b2_dir, b2_docx)

    # 3) Match
    raw_csv = os.path.join(workdir, f"matches_{book1_id}_vs_{book2_id}.csv")
    match_books_sequential(b1_docx, b2_docx, raw_csv, min_len_chars=min_len_chars, score_threshold=score_threshold)

    # 4) Add page numbers by searching per-page TXT
    final_csv = os.path.join(workdir, f"matches_{book1_id}_vs_{book2_id}_with_pages.csv")
    update_csv_with_pages(
        raw_csv, final_csv,
        b1_pages_dir=b1_dir, b2_pages_dir=b2_dir,
        book1_use_fuzzy=use_fuzzy_backmap, book2_use_fuzzy=use_fuzzy_backmap,
        cols=("pos_in_book1","pos_in_book2","sentence_book1","sentence_book2")  # columns we wrote above
    )

    return {
        "book1_txt_dir": b1_dir,
        "book2_txt_dir": b2_dir,
        "book1_docx": b1_docx,
        "book2_docx": b2_docx,
        "raw_csv": raw_csv,
        "final_csv": final_csv
    }
